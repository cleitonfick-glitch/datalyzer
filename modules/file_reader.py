"""
modules/file_reader.py
======================
Leitura e extração de dados de diferentes formatos de arquivo.
Suporta: Excel (.xlsx/.xls), CSV, PDF e Word (.docx).
"""

import pandas as pd
import io
import re
from typing import Optional


class FileReader:
    """
    Centraliza a leitura de arquivos nos formatos suportados.
    Detecta o tipo pelo nome/extensão e delega ao método correto.
    """

    def read(self, uploaded_file) -> Optional[pd.DataFrame]:
        """
        Ponto de entrada único: recebe um UploadedFile do Streamlit
        e retorna um DataFrame pandas.

        Args:
            uploaded_file: objeto streamlit.UploadedFile

        Returns:
            pd.DataFrame ou None se não houver dados extraíveis

        Raises:
            ValueError: se o formato não for suportado
            Exception: erros de leitura específicos por formato
        """
        name = uploaded_file.name.lower()
        ext = name.rsplit(".", 1)[-1]

        if ext in ("xlsx", "xls"):
            return self._read_excel(uploaded_file)
        elif ext == "csv":
            return self._read_csv(uploaded_file)
        elif ext == "pdf":
            return self._read_pdf(uploaded_file)
        elif ext in ("docx", "doc"):
            return self._read_word(uploaded_file)
        else:
            raise ValueError(
                f"Formato '.{ext}' não suportado. "
                "Use: xlsx, xls, csv, pdf, docx"
            )

    # ─────────────────────────────────────────
    # Excel
    # ─────────────────────────────────────────
    def _read_excel(self, uploaded_file) -> pd.DataFrame:
        """
        Lê arquivo Excel (.xlsx ou .xls).
        Se houver múltiplas abas, concatena todas ou usa a primeira
        conforme configuração.
        """
        raw = uploaded_file.read()
        xls = pd.ExcelFile(io.BytesIO(raw))
        sheet_names = xls.sheet_names

        if len(sheet_names) == 1:
            df = pd.read_excel(xls, sheet_name=0)
        else:
            # Concatena todas as abas com coluna identificadora
            frames = []
            for sheet in sheet_names:
                try:
                    tmp = pd.read_excel(xls, sheet_name=sheet)
                    if not tmp.empty:
                        tmp["_aba"] = sheet
                        frames.append(tmp)
                except Exception:
                    pass
            if not frames:
                raise ValueError("Nenhuma aba com dados encontrada.")
            df = pd.concat(frames, ignore_index=True)

        return self._infer_dtypes(df)

    # ─────────────────────────────────────────
    # CSV
    # ─────────────────────────────────────────
    def _read_csv(self, uploaded_file) -> pd.DataFrame:
        """
        Lê CSV detectando automaticamente separador e encoding.
        Tenta: utf-8, latin-1, cp1252.
        Detecta separadores: vírgula, ponto-e-vírgula, tab e pipe.
        """
        raw = uploaded_file.read()
        encodings = ["utf-8", "latin-1", "cp1252"]
        separators = [",", ";", "\t", "|"]

        for enc in encodings:
            try:
                text = raw.decode(enc)
                # Detecta separador pela primeira linha
                first_line = text.split("\n")[0]
                sep = max(separators, key=lambda s: first_line.count(s))
                df = pd.read_csv(
                    io.StringIO(text),
                    sep=sep,
                    on_bad_lines="skip"
                )
                if df.shape[1] > 1:  # Separador correto encontrado
                    return self._infer_dtypes(df)
            except UnicodeDecodeError:
                continue

        raise ValueError(
            "Não foi possível decodificar o CSV. "
            "Verifique o encoding do arquivo."
        )

    # ─────────────────────────────────────────
    # PDF
    # ─────────────────────────────────────────
    def _read_pdf(self, uploaded_file) -> pd.DataFrame:
        """
        Extrai tabelas e texto de PDFs usando pdfplumber (preferido)
        com fallback para PyPDF2. Retorna a maior tabela encontrada
        ou um DataFrame com o texto por página.
        """
        raw = uploaded_file.read()

        # Tentativa 1: pdfplumber (extração de tabelas)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                all_tables = []
                all_text = []

                for page in pdf.pages:
                    # Extrair tabelas
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            all_tables.append(table)

                    # Extrair texto como fallback
                    text = page.extract_text()
                    if text:
                        all_text.append(text)

                if all_tables:
                    # Usa a maior tabela encontrada
                    biggest = max(all_tables, key=lambda t: len(t))
                    if len(biggest) > 1:
                        df = pd.DataFrame(biggest[1:], columns=biggest[0])
                        return self._infer_dtypes(df)

                # Sem tabelas: estrutura o texto
                if all_text:
                    return self._text_to_dataframe(all_text)

        except ImportError:
            pass  # pdfplumber não instalado, tenta PyPDF2

        # Tentativa 2: PyPDF2 (texto simples)
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)

            if texts:
                return self._text_to_dataframe(texts)

        except ImportError:
            raise ImportError(
                "Instale pdfplumber ou PyPDF2: pip install pdfplumber PyPDF2"
            )

        raise ValueError("Nenhum dado extraível encontrado no PDF.")

    def _text_to_dataframe(self, texts: list) -> pd.DataFrame:
        """
        Converte lista de textos de páginas em DataFrame estruturado.
        Detecta padrões tabulares separados por espaços ou delimitadores.
        """
        rows = []
        for i, text in enumerate(texts):
            lines = [l.strip() for l in text.split("\n") if l.strip()]
            for line in lines:
                rows.append({"pagina": i + 1, "conteudo": line})

        if not rows:
            raise ValueError("PDF não contém texto extraível.")

        df = pd.DataFrame(rows)

        # Tenta detectar colunas em texto delimitado
        sample = "\n".join([r["conteudo"] for r in rows[:20]])
        for sep in ["\t", "  ", " | "]:
            if sep in sample:
                structured_rows = []
                for r in rows:
                    parts = [p.strip() for p in r["conteudo"].split(sep)
                             if p.strip()]
                    if len(parts) > 1:
                        structured_rows.append(parts)
                if structured_rows:
                    max_cols = max(len(r) for r in structured_rows)
                    df2 = pd.DataFrame(
                        structured_rows,
                        columns=[f"col_{i+1}" for i in range(max_cols)]
                    )
                    return self._infer_dtypes(df2)

        return df

    # ─────────────────────────────────────────
    # Word
    # ─────────────────────────────────────────
    def _read_word(self, uploaded_file) -> pd.DataFrame:
        """
        Extrai tabelas de documentos Word (.docx).
        Se não houver tabelas, retorna parágrafos estruturados.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "Instale python-docx: pip install python-docx"
            )

        raw = uploaded_file.read()
        doc = Document(io.BytesIO(raw))

        # Extrair tabelas
        all_tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                all_tables.append(rows)

        if all_tables:
            # Usa a maior tabela
            biggest = max(all_tables, key=len)
            if len(biggest) > 1:
                header = biggest[0]
                # Garante colunas únicas
                seen = {}
                clean_header = []
                for h in header:
                    if h in seen:
                        seen[h] += 1
                        clean_header.append(f"{h}_{seen[h]}")
                    else:
                        seen[h] = 0
                        clean_header.append(h if h else f"col_{len(clean_header)+1}")

                df = pd.DataFrame(biggest[1:], columns=clean_header)
                return self._infer_dtypes(df.infer_objects())

        # Fallback: parágrafos
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("Documento Word não contém dados extraíveis.")

        df = pd.DataFrame({
            "linha": range(1, len(paragraphs) + 1),
            "conteudo": paragraphs
        })
        return df

    # ─────────────────────────────────────────
    # Utilitário: inferência de tipos
    # ─────────────────────────────────────────
    def _infer_dtypes(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Tenta converter colunas string para tipos mais específicos:
        numérico → datetime → mantém string.
        Remove espaços em branco dos nomes de coluna.
        """
        # Limpa nomes de colunas
        df.columns = [
            str(c).strip().replace("\n", " ") for c in df.columns
        ]

        for col in df.columns:
            if df[col].dtype != object:
                continue

            # Limpa strings
            df[col] = df[col].astype(str).str.strip()
            df[col] = df[col].replace({"nan": None, "None": None, "": None})

            # Tenta numérico
            try:
                numeric = pd.to_numeric(
                    df[col].str.replace(",", ".", regex=False)
                               .str.replace(r"[^\d.\-]", "", regex=True),
                    errors="coerce"
                )
                non_null = numeric.notna().sum()
                if non_null / max(len(df), 1) > 0.5:  # >50% conversível
                    df[col] = numeric
                    continue
            except Exception:
                pass

            # Tenta data
            try:
                dates = pd.to_datetime(df[col], errors="coerce", dayfirst=True)
                non_null = dates.notna().sum()
                if non_null / max(len(df), 1) > 0.5:
                    df[col] = dates
                    continue
            except Exception:
                pass

        return df
